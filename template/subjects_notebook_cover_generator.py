from copy import deepcopy

from docx import Document

document = Document(input('Название файла с шаблоном (включая расширение .docx/.doc): '))

peoples = ['Иванова Ивана', 'Петрова Петра', 'Глазунова Никиты', 'TheCode Factory']

klass = input('Класс (без литеры): ')
litera = input('Литера: ')
school = input('Название школы: ')
is_big_subject_name = int(input('Название предмета состоит из двух строк? (1 - да, 0 - нет): '))

index = 0

while index != len(peoples) - 1:
    document.paragraphs[0]._p.addnext(deepcopy(document.tables[0]._tbl))
    index += 1

index = 0

klass_litera_index = 4
school_index = 5
name_surname_index = 6

if is_big_subject_name == 1:
    klass_litera_index += 1
    school_index += 1
    name_surname_index += 1

for table in document.tables:
    table.rows[0].cells[1].paragraphs[klass_litera_index].text = f"ученика(-цы) {klass} «{litera}» класса"
    table.rows[0].cells[1].paragraphs[school_index].text = school
    table.rows[0].cells[1].paragraphs[name_surname_index].text = peoples[index]
    table.rows[0].cells[1].paragraphs[name_surname_index].style = document.styles['No Spacing']

    print(f'№{index}')
    for paragraph in table.rows[0].cells[1].paragraphs:
        if paragraph.text != '':
            print(paragraph.text)
    print('\n==========\n')

    index += 1

safe_file = input("Название файла для сохранения (без расширения): ")
document.save(f'success/{safe_file}.docx')
print(f'Файл сохранён в папке success как {safe_file}.docx.')